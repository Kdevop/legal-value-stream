
"""
Contract Processor - Production Script
Parses legal contracts (PDF/DOCX) and extracts risk-flagged clauses using LLM analysis.
"""

import os
import sys
import json
import argparse
from pathlib import Path
from typing import List, Optional

import pdfplumber
from docx import Document
import openai
from pydantic import BaseModel, Field
from dotenv import load_dotenv

# for testing the logger before implementation
#from utils.audit import setup_audit_logger, log_ai_interaction

load_dotenv(override=True)
OPENROUTER_API_KEY = os.getenv('OPENROUTER_API_KEY')

# ============================================================
# PYDANTIC MODELS
# ============================================================

class ClauseModel(BaseModel):
    clause_type: str = Field(description="e.g., Liability, Data Training, Audit Rights")
    text: str = Field(description="The verbatim text extracted from the contract")
    page_number: int = Field(description="The page where this clause was found")
    risk_flag: str = Field(description="RED, YELLOW, or GREEN")
    risk_explanation: str = Field(description="Reasoning for the risk flag based on 2026 standards")
    risk_score: int = Field(ge=1, le=10, description="1-10 priority score")


class ContractAnalysis(BaseModel):
    clauses: List[ClauseModel]


# ============================================================
# DOCUMENT PARSER
# ============================================================

class DocParser:
    """
    Parses contract documents (PDF and DOCX) into page-indexed dictionaries.
    """

    def __init__(self):
        self.supported_formats = ['.pdf', '.docx']

    def parse_file(self, file_path: str) -> List[dict]:
        """
        Determines file type and routes to the appropriate parser.
        Returns: List of dicts, e.g., [{'page': 1, 'text': '...'}, {'page': 2, 'text': '...'}]
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f'File {file_path} not found')

        if file_path.suffix.lower() not in self.supported_formats:
            raise ValueError(f"Unsupported file format {file_path.suffix}")

        if file_path.suffix.lower() == '.pdf':
            return self._parse_pdf(file_path)
        elif file_path.suffix.lower() == '.docx':
            return self._parse_docx(file_path)

    def _parse_pdf(self, file_path: Path) -> List[dict]:
        """
        Extracts text page-by-page from PDF.
        """
        pages_data = []
        try:
            with pdfplumber.open(file_path) as pdf:
                print(f"Parsing PDF: {file_path.name} ({len(pdf.pages)} pages)")
                for page_num, page in enumerate(pdf.pages, start=1):
                    page_text = page.extract_text()
                    if page_text:
                        pages_data.append({
                            "page": page_num,
                            "text": page_text
                        })
            
            total_chars = sum(len(p['text']) for p in pages_data)
            print(f"✅ Extracted {total_chars} characters across {len(pages_data)} pages")
            return pages_data
        except Exception as e:
            raise RuntimeError(f"Error parsing PDF {file_path}: {e}")

    def _parse_docx(self, file_path: Path) -> List[dict]:
        """
        Extracts text from DOCX.
        """
        try:
            doc = Document(file_path)
            print(f"Parsing DOCX: {file_path.name}")

            text_cont = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_cont.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_cont.append(cell.text)
            
            full_text = "\n\n".join(text_cont)
            
            return [{"page": 1, "text": full_text}]
            
        except Exception as e:
            raise RuntimeError(f"Error parsing DOCX {file_path}: {e}")


# ============================================================
# CLAUSE EXTRACTOR
# ============================================================

class ClauseExtractor:
    """
    Extracts and categorises contract clauses using LLM.
    Uses strict 2026 Risk Rubrics and forces structured JSON output.
    """

    def __init__(self, api_key: str):
        self.client = openai.OpenAI(
            base_url="https://openrouter.ai/api/v1",
            api_key=api_key
        )
        self.model = "google/gemini-2.5-flash-lite"
        print(f"Initializing Clause Extractor with model: {self.model}")

    def _get_system_prompt(self) -> str:
        """
        Returns the detailed legal logic for the AI to follow.
        """
        return """You are a senior Legal Contract Analyst specialising in UK commercial contracts and 2026 regulatory compliance.

Your task is to analyse a page of a legal document (contract or court judgment) and extract clauses relevant to risk assessment.

Return ONLY valid JSON that matches this exact structure:

{
  "clauses": [
    {
      "clause_type": "string",
      "text": "verbatim clause text from the document",
      "page_number": integer,
      "risk_flag": "RED | YELLOW | GREEN",
      "risk_explanation": "short explanation of the risk",
      "risk_score": integer (1-10)
    }
  ]
}

Rules:
- Use the exact field names shown above.
- Quote clause text verbatim from the document.
- Do not invent clauses.
- If no relevant clauses are present return: {"clauses": []}
- Do not include any text outside the JSON.

Risk Rubric:

LIABILITY
RED: unlimited liability or missing liability caps  
YELLOW: liability cap present but less than 2x contract value  
GREEN: reasonable liability cap (≈2x contract value or higher)

DATA TRAINING
RED: vendor can use customer data/IP to train AI models  
GREEN: explicit prohibition or opt-out for model training

AUDIT RIGHTS
RED: no audit rights or transparency for AI systems  
GREEN: customer has audit rights or model transparency

TERMINATION
RED: no termination or exit rights  
YELLOW: termination allowed but heavy restrictions (>90 day notice or large fees)  
GREEN: reasonable termination for convenience

GOVERNING LAW
RED: jurisdiction outside the UK for a UK entity  
GREEN: UK governing law (England & Wales, Scotland, or NI)

Risk Score Guidance (1–10):
1–3: low concern  
4–6: moderate concern  
7–8: high concern  
9–10: critical contractual risk"""

    def extract_clauses(self, pages: List[dict]) -> dict:
        # Set up the audit logger for testing, the logger set up might need moving for production
        #  logger = setup_audit_logger()

        """
        Processes each page and aggregates extracted clauses.
        :param pages: List of dicts [{'page': 1, 'text': '...'}, ...]
        """
        all_extracted_clauses = []
        
        print(f"Extracting clauses from {len(pages)} pages...")

        for page_data in pages:
            page_num = page_data['page']
            text = page_data['text']

            try:
                response = self.client.chat.completions.create(
                    model=self.model,
                    messages=[
                        {"role": "system", "content": self._get_system_prompt()},
                        {"role": "user", "content": f"ANALYZE PAGE {page_num}:\n\n{text}"}
                    ],
                    response_format={"type": "json_object"},
                    temperature=0.1
                )

                response_content = response.choices[0].message.content
                page_results = json.loads(response_content)
                # log_ai_interaction(logger, user_input=f"ANALYZE PAGE {page_num}:\n\n{text[:500]}...",  # Log first 500 chars for brevity
                #                    ai_output=response_content,
                #                    model_name=self.model,
                #                    metadata={"page_number": page_num})

                for clause in page_results.get("clauses", []):
                    clause["page_found"] = page_num
                    
                    validated_clause = ClauseModel(**clause)
                    all_extracted_clauses.append(validated_clause.model_dump())

            except Exception as e:
                print(f"⚠️  Error processing Page {page_num}: {e}", file=sys.stderr)

        return {"clauses": all_extracted_clauses}


# ============================================================
# MAIN EXECUTION
# ============================================================

def print_audit_summary(clauses: List[dict]):
    """
    Prints a formatted audit summary to console.
    """
    print(f"\n{'='*60}")
    print(f"AUDIT SUMMARY: {len(clauses)} CLAUSES IDENTIFIED")
    print(f"{'='*60}")

    for i, clause in enumerate(clauses, 1):
        flag = clause['risk_flag']
        color = "🔴" if flag == "RED" else "🟡" if flag == "YELLOW" else "🟢"
        
        print(f"\n{i}. {color} [{flag}] Type: {clause['clause_type']}")
        print(f"   📍 Location: Page {clause['page_number']}")
        print(f"   ⚖️  Risk Score: {clause['risk_score']}/10")
        print(f"   📝 Verbatim: \"{clause['text'][:200]}...\"")
        print(f"   💡 Analysis: {clause['risk_explanation']}")
        print(f"{'-'*60}")


def main():
    parser = argparse.ArgumentParser(
        description="Contract Processor - Extract and analyze contract clauses"
    )
    parser.add_argument(
        "contract_path",
        type=str,
        help="Path to the contract file (PDF or DOCX)"
    )
    parser.add_argument(
        "-o", "--output",
        type=str,
        default="contract_audit_results.json",
        help="Output JSON file path (default: contract_audit_results.json)"
    )
    
    args = parser.parse_args()

    # Load environment variables
    load_dotenv()
    api_key = os.getenv('OPENROUTER_API_KEY')
    
    if not api_key:
        print("❌ Error: OPENROUTER_API_KEY not found in environment", file=sys.stderr)
        print("Please create a .env file with your API key", file=sys.stderr)
        sys.exit(1)

    try:
        # Parse the document
        doc_parser = DocParser()
        contract_pages = doc_parser.parse_file(args.contract_path)

        if not contract_pages:
            print("❌ No text extracted. Please check the file path and format.", file=sys.stderr)
            sys.exit(1)

        # Extract clauses
        extractor = ClauseExtractor(api_key=api_key)
        analysis_result = extractor.extract_clauses(contract_pages)
        
        extracted_clauses = analysis_result.get("clauses", [])

        if not extracted_clauses:
            print("⚠️  No clauses extracted from the document")
            sys.exit(0)

        # Display results
        print_audit_summary(extracted_clauses)

        # Save to JSON
        with open(args.output, "w") as f:
            json.dump(extracted_clauses, f, indent=4)
        
        print(f"\n✅ Audit complete. Data saved to '{args.output}'")

    except FileNotFoundError as e:
        print(f"❌ Error: {e}", file=sys.stderr)
        sys.exit(1)
    except ValueError as e:
        print(f"❌ Error: {e}", file=sys.stderr)
        sys.exit(1)
    except RuntimeError as e:
        print(f"❌ Error: {e}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"❌ Unexpected error: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
